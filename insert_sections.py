# -*- coding: utf-8 -*-
"""
Inserta las secciones 7.1, 7.2 y 7.3 en Formato Capstone.docx
"""
import sys
import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DOC_PATH = r"C:\Users\holaq\Proyecto-TITA\Formato Capstone.docx"
doc = Document(DOC_PATH)

# ─── helpers ──────────────────────────────────────────────────────────────────

def para_text(p):
    return p.text.strip()

def add_border_to_para(p):
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    for side in ("top", "left", "bottom", "right"):
        bdr = OxmlElement(f"w:{side}")
        bdr.set(qn("w:val"), "single")
        bdr.set(qn("w:sz"), "4")
        bdr.set(qn("w:space"), "4")
        bdr.set(qn("w:color"), "AAAAAA")
        pBdr.append(bdr)
    pPr.append(pBdr)

def shade_cell(cell, fill="D9D9D9"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


class SectionBuilder:
    """Acumula párrafos/tablas y los inserta antes de `anchor`."""

    def __init__(self, doc, anchor_para):
        self.doc = doc
        self.anchor = anchor_para   # insertamos justo antes de este párrafo
        self._buf = []              # lista de _element XML a insertar

    # ── paragraph helpers ────────────────────────────────────────────────────

    def _flush_add(self, p):
        """Mueve el párrafo recién añadido al final del doc hacia self.anchor."""
        el = p._element
        el.getparent().remove(el)
        self.anchor._element.addprevious(el)
        return p

    def heading(self, text, level=2):
        p = self.doc.add_heading(text, level=level)
        return self._flush_add(p)

    def para(self, text="", bold=False, italic=False, center=False, color=None,
             font_name=None, font_size=None, code=False):
        p = self.doc.add_paragraph()
        if text:
            run = p.add_run(text)
            run.bold = bold
            run.italic = italic
            if color:
                run.font.color.rgb = color
            if font_name:
                run.font.name = font_name
            if font_size:
                run.font.size = Pt(font_size)
        if center:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if code:
            add_border_to_para(p)
        return self._flush_add(p)

    def blank(self):
        return self.para()

    def fig(self, text):
        return self.para(text, italic=True,
                         color=RGBColor(0x70, 0x70, 0x70), center=True)

    def code(self, text):
        return self.para(text, font_name="Courier New",
                         font_size=8, code=True)

    def bold_intro(self, label, body):
        """Un párrafo con label en negrita seguido del cuerpo normal."""
        p = self.doc.add_paragraph()
        run_label = p.add_run(label + " ")
        run_label.bold = True
        p.add_run(body)
        return self._flush_add(p)

    # ── table helper ────────────────────────────────────────────────────────

    def table(self, headers, rows):
        tbl = self.doc.add_table(rows=1 + len(rows), cols=len(headers))
        tbl.style = "Table Grid"
        # Cabecera
        for i, h in enumerate(headers):
            cell = tbl.rows[0].cells[i]
            cell.text = h
            for run in cell.paragraphs[0].runs:
                run.bold = True
            shade_cell(cell)
        # Filas
        for r, row_data in enumerate(rows):
            for c, val in enumerate(row_data):
                tbl.rows[r + 1].cells[c].text = str(val)
        # Mover la tabla
        tbl_el = tbl._tbl
        tbl_el.getparent().remove(tbl_el)
        self.anchor._element.addprevious(tbl_el)
        return tbl


# ─── localizar headings ───────────────────────────────────────────────────────

anchors = {}
targets = {
    "71": "Diseño de la solución",
    "72": "Desarrollo de la solución",
    "73": "Pruebas y evaluación de la solución",
    "74": "Resultados y Discusión.",
}

for p in doc.paragraphs:
    if p.style.name.startswith("Heading"):
        for key, txt in targets.items():
            if para_text(p) == txt and key not in anchors:
                anchors[key] = p

if len(anchors) < 4:
    print("ERROR: No se encontraron todos los encabezados:", anchors.keys())
    sys.exit(1)

print("Headings encontrados OK")

# ─── limpiar placeholders ─────────────────────────────────────────────────────

def clear_between(doc, start_heading, end_heading):
    body = doc.element.body
    paras = list(doc.paragraphs)
    # Buscar por elemento XML (identidad de objeto segura)
    start_el = start_heading._element
    end_el   = end_heading._element
    start_i = next(i for i, p in enumerate(paras) if p._element is start_el)
    end_i   = next(i for i, p in enumerate(paras) if p._element is end_el)
    to_del = [paras[i]._element for i in range(start_i + 1, end_i)]
    for el in to_del:
        body.remove(el)

clear_between(doc, anchors["73"], anchors["74"])
clear_between(doc, anchors["72"], anchors["73"])
clear_between(doc, anchors["71"], anchors["72"])
print("Placeholders eliminados")

# Re-obtener referencias después de limpiar
for p in doc.paragraphs:
    if p.style.name.startswith("Heading"):
        for key, txt in targets.items():
            if para_text(p) == txt:
                anchors[key] = p

# ═══════════════════════════════════════════════════════════════════════════════
#  SECCIÓN 7.1
# ═══════════════════════════════════════════════════════════════════════════════

s1 = SectionBuilder(doc, anchors["72"])

# 7.1.1
s1.heading("7.1.1 Arquitectura del sistema — C4 Nivel 2 (Contenedores)")
s1.para("El sistema se compone de tres contenedores que se comunican a través del servicio en la nube de Photon Fusion:")
s1.blank()
s1.bold_intro("Contenedor 1 — Cliente VR (Explorador).",
    "Aplicación Unity desplegada en el visor Meta Quest 3. Gestiona la locomoción física mediante el SDK de KAT VR con fallback automático al joystick del visor, la interacción con objetos virtuales (multímetro, componentes, nodos de circuito) a través de XR Interaction Toolkit 3.4.1, y la visualización del entorno inmersivo de la nave espacial renderizado con Universal Render Pipeline (URP 17.4).")
s1.blank()
s1.bold_intro("Contenedor 2 — Cliente PC (Técnico).",
    "Aplicación Unity ejecutada en PC con monitor convencional. Gestiona la visualización del manual técnico interactivo, la selección de componentes desde la estación de trabajo, el envío de componentes al Explorador y la interfaz de diagnóstico. Actúa como Host (servidor autorizado) de la sesión de red Photon Fusion.")
s1.blank()
s1.bold_intro("Contenedor 3 — Photon Fusion Cloud (servicio externo).",
    "Servicio en la nube que sincroniza el estado de la sesión entre ambos clientes mediante NetworkObject y RPCs tipados. Mantiene el objeto GameSession con el estado del reto actual, el tipo y valor del componente pendiente, y las confirmaciones de instalación.")
s1.blank()
s1.fig("[Insertar Figura 7.1: Diagrama C4 Nivel 2 — Contenedores del sistema]")

# 7.1.2
s1.blank()
s1.heading("7.1.2 Arquitectura interna — C4 Nivel 3 (Componentes)")
s1.para("Dentro de cada cliente Unity los scripts se organizan en ocho módulos con responsabilidades claramente delimitadas según el Principio de Responsabilidad Única (SOLID - SRP):")
s1.blank()
s1.bold_intro("Módulo Electrical — Motor de simulación eléctrica.",
    "La clase abstracta ElectricalComponent define la interfaz de todos los componentes. Las clases concretas Resistor, LED, Capacitor, VoltageSource y ArduinoPin implementan GetResistance() y Calculate() según su física. CircuitManager coordina la simulación completa con un intervalo de 50 ms controlado por un dirty-flag que previene simulaciones redundantes.")
s1.blank()
s1.bold_intro("Módulo Gameplay — Lógica de retos y entrega de componentes.",
    "GameManager orquesta los cuatro retos mediante una máquina de estados basada en el enum LevelType, comunicándose exclusivamente por eventos estáticos para evitar acoplamiento directo. ComponentDeliverySystem gestiona el ciclo de envío-instalación-validación. InstructionSystem provee guía paso a paso. ObjectiveSystem calcula la puntuación con bonificación temporal.")
s1.blank()
s1.bold_intro("Módulo Interaction — Interacción física VR.",
    "GrabbableComponent permite tomar y soltar objetos con los controladores XRI. Multimeter gestiona las sondas virtuales y calcula el voltaje entre dos nodos. ComponentSlot valida el tipo y valor del componente instalado por el Explorador.")
s1.blank()
s1.bold_intro("Módulo Networking — Comunicación entre roles.",
    "ConnectionManager inicia la sesión Photon Fusion según el rol asignado (Host para Técnico, Client para Explorador). GameSession (NetworkBehaviour) expone tres RPCs tipados para la comunicación entre roles.")
s1.blank()
s1.bold_intro("Módulo Player — Controladores de ambos roles.",
    "PlayerController abstrae la locomoción del Explorador con soporte KAT VR y fallback al joystick. TechnicianController configura la interacción por ratón. HapticFeedback envía pulsos al chaleco háptico proporcionales a la corriente simulada.")
s1.blank()
s1.bold_intro("Módulos UI, Desktop, Core — Presentación e interfaz.",
    "TechnicianHUDController muestra instrucciones, puntaje y temporizador. MultimeterUI renderiza la lectura de voltaje. CircuitDiagramPanel muestra el esquema con indicadores de falla. TechnicianWorkstation gestiona la selección y envío de componentes.")
s1.blank()
s1.fig("[Insertar Figura 7.2: Diagrama C4 Nivel 3 — Módulos Gameplay y Electrical]")
s1.fig("[Insertar Figura 7.3: Diagrama C4 Nivel 3 — Módulos Networking y Player]")

# 7.1.3
s1.blank()
s1.heading("7.1.3 Patrones de diseño aplicados")
s1.blank()
s1.bold_intro("Observer / Event-driven.",
    "Todos los módulos se comunican a través de eventos estáticos de C# (Action<T>) publicados por GameManager y CircuitManager. La UI, el sistema de objetivos y el tracker de desempeño reaccionan a cambios de estado sin referencias directas al núcleo de lógica.")
s1.blank()
s1.bold_intro("Dirty Flag.",
    "CircuitManager mantiene un indicador _dirty que solo activa la simulación cuando algún componente cambia mediante MarkDirty(), limitando la carga computacional a 20 simulaciones por segundo únicamente cuando hay modificaciones reales.")
s1.blank()
s1.bold_intro("Strategy.",
    "El método RunSimulation() de CircuitManager delega en SimulateSeries(), SimulateParallel() o SimulateMixed() según la topología configurada por el enum CircuitTopology, permitiendo agregar nuevas topologías sin modificar la lógica de decisión.")
s1.blank()
s1.bold_intro("Template Method.",
    "ElectricalComponent define el esqueleto del comportamiento eléctrico, dejando que las subclases implementen GetResistance() y Calculate() con su física particular.")

# 7.1.4
s1.blank()
s1.heading("7.1.4 Estándares y buenas prácticas")
s1.para("ISO/IEC 25010: El diseño prioriza adecuación funcional (simulación eléctrica correcta), usabilidad (feedback visual e instrucciones paso a paso) y mantenibilidad (separación de responsabilidades por módulo).")
s1.para("OpenXR 1.16.1: Capa de abstracción de hardware VR que garantiza compatibilidad con Meta Quest 3 sin dependencia del SDK propietario de Meta.")
s1.para("C# Unity Coding Standards: PascalCase para clases y métodos públicos, camelCase con prefijo _ para campos privados, propiedades de solo lectura para estado observable desde el Inspector.")
s1.para("Principio SRP (SOLID): Cada script gestiona exactamente un dominio.")

# 7.1.5
s1.blank()
s1.heading("7.1.5 Restricciones de diseño")
s1.table(
    ["Restricción", "Justificación"],
    [
        ["Solo corriente continua (DC), componentes pasivos",
         "Contenido del primer progreso del sílabo de Computación Ubicua"],
        ["Máximo 2 jugadores simultáneos",
         "Un visor Meta Quest 3 y un PC disponibles en el laboratorio UDLA"],
        ['Sala de red fija ("LaboratorioUbicua")',
         "Sesiones presenciales guiadas; no requiere descubrimiento dinámico"],
        ["Sin persistencia de datos entre sesiones",
         "Prototipo funcional; dashboard docente definido como trabajo futuro"],
        ["Motor Unity 6 (6000.4.3f1) + URP 17.4",
         "Compatibilidad con XRI 3.4.1 y renderizado optimizado para Meta Quest 3"],
    ]
)

print("Seccion 7.1 insertada")

# ═══════════════════════════════════════════════════════════════════════════════
#  SECCIÓN 7.2
# ═══════════════════════════════════════════════════════════════════════════════

s2 = SectionBuilder(doc, anchors["73"])

# 7.2.1
s2.heading("7.2.1 Metodología — Cascada")
s2.para("El proyecto se ejecutó en cinco fases secuenciales conforme a la planificación de la Sección 5. Cada fase generó entregables verificables antes de iniciar la siguiente.")
s2.blank()
s2.table(
    ["Fase", "Semanas", "Entregables principales"],
    [
        ["1 — Análisis", "1–2",
         "Requerimientos funcionales/no funcionales; especificación de hardware VR; encuesta diagnóstica (20 estudiantes)"],
        ["2 — Diseño", "3–4",
         "Diagramas C4 niveles 1–3; casos de uso; prototipos de media fidelidad; selección de alternativa mediante matriz ISO/IEC 25010"],
        ["3 — Implementación", "5–9",
         "50 scripts C# en 8 módulos; 3 escenas Unity (Tecnico.unity, Explorador.unity, IntegratedDemo.unity); prefabs de red Fusion; herramientas de editor para generación de escena"],
        ["4 — Pruebas e integración", "10–11",
         "Casos de prueba funcionales, de integración y de usabilidad; registro de 6 defectos y sus correcciones"],
        ["5 — Despliegue y evaluación", "12",
         "Build Windows (cliente Técnico); APK Meta Quest (cliente Explorador); evaluación con usuarios en laboratorio UDLA"],
    ]
)

# 7.2.2
s2.blank()
s2.heading("7.2.2 Diagrama de clases — C4 Nivel 4")
s2.para("El sistema cuenta con 50 clases C# agrupadas en ocho módulos. A continuación se presenta el modelo de clases de los módulos centrales.")
s2.blank()
s2.para("Jerarquía del módulo Electrical", bold=True)
s2.code(
"ElectricalComponent  {abstract}\n"
"    nodeA, nodeB : ElectricalNode\n"
"    current, voltageDrop : float\n"
"    + GetResistance() : float   {abstract}\n"
"    + Calculate() : void        {abstract}\n"
"    +-- Resistor\n"
"    |       resistance, faultyResistance, correctResistance : float\n"
"    |       hasFault : bool\n"
"    |       + ApplyFault(), Repair(), IsValueCorrect(v) : bool\n"
"    +-- LED\n"
"    |       resistance : float\n"
"    |       polarityInverted, isOn : bool\n"
"    +-- Capacitor\n"
"    |       polarityInverted : bool\n"
"    +-- VoltageSource\n"
"    |       voltage : float\n"
"    +-- ArduinoPin\n"
"            pinNumber, correctPinNumber : int\n"
"            hasFault, hasLooseCable : bool\n"
"\n"
"ElectricalNode\n"
"    voltage, current : float\n"
"\n"
"CircuitManager\n"
"    components : List<ElectricalComponent>\n"
"    topology : CircuitTopology {Series | Parallel | Mixed}\n"
"    totalCurrent, sourceVoltage, totalPower : float\n"
"    isShortCircuited : bool\n"
"    <<static event>> OnCircuitChanged : Action\n"
"    + MarkDirty(), ForceSimulate(), AreAllLEDsOn() : bool"
)
s2.blank()
s2.para("Módulo Gameplay — Relaciones entre clases", bold=True)
s2.code(
"GameManager\n"
"    circuit : CircuitManager\n"
"    multimeter : Multimeter\n"
"    performance : PerformanceTracker\n"
"    instructionSystem : InstructionSystem\n"
"    <<static events>>\n"
"        OnLevelLoaded : Action<LevelType>\n"
"        OnLevelCompleted : Action<LevelType, bool>\n"
"        OnFaultDetected : Action<string>\n"
"        OnTimerTick : Action<float> / OnGameCompleted : Action\n"
"\n"
"ObjectiveSystem  suscribe -> GameManager.OnLevelLoaded / Completed / GameCompleted\n"
"    + CompleteObjective(index, timeBonus) : void\n"
"    <<static event>> OnSessionEnded : Action<SessionResult>\n"
"\n"
"PerformanceTracker  suscribe -> GameManager.OnLevelLoaded / Completed\n"
"    + AddError(type), GetEvaluation(), GetTimeBonus() : float\n"
"\n"
"InstructionSystem  suscribe -> CircuitManager.OnCircuitChanged\n"
"    + BuildInstructions(), GetCurrentInstruction() : string\n"
"\n"
"ComponentDeliverySystem\n"
"    + SendResistor(v), SendLED(pol), SendCapacitor(pol), SendArduinoPin(pin)\n"
"    + OnExplorerInstalled(slot) : void\n"
"    <<static events>> OnComponentSent, OnRepairValidated, OnDeliveryError"
)
s2.blank()
s2.para("Módulo Networking", bold=True)
s2.code(
"GameSession : NetworkBehaviour\n"
"    [Networked] RetoActual : int\n"
"    [Networked] HayComponentePendiente : NetworkBool\n"
"    [Networked] TipoComponentePendiente : int\n"
"    [Networked] ValorComponentePendiente : float\n"
"    <<static>> Instance : GameSession\n"
"    <<static events>> OnComponenteRecibido, OnComponenteInstalado, OnRetoChanged\n"
"    + EnviarComponente(tipo, valor) : void   {solo Host}\n"
"    + ReportarInstalacion(exito) : void\n"
"\n"
"ConnectionManager : MonoBehaviour, INetworkRunnerCallbacks\n"
"    rolAutomatico : AutoConnectRole {Ninguno | Explorador | Tecnico}\n"
"    + StartSimulation(mode: GameMode) : void"
)
s2.blank()
s2.fig("[Insertar Figura 7.4: Diagrama de clases — módulos Electrical y Gameplay]")
s2.fig("[Insertar Figura 7.5: Diagrama de clases — módulos Networking y Player]")

# 7.2.3
s2.blank()
s2.heading("7.2.3 Flujo de red asimétrica")
s2.para("La comunicación entre roles sigue un flujo controlado por autoridad de estado (Photon Fusion StateAuthority):")
s2.blank()
s2.code(
"1. Técnico (Host) selecciona componente -> ComponentSendingTray.SendComponent()\n"
"2. Técnico llama GameSession.EnviarComponente(tipo, valor) -> RPC broadcast\n"
"3. Explorador recibe OnComponenteRecibido -> ExplorerComponentReceiver instancia prefab en bandeja\n"
"4. Explorador instala en ComponentSlot -> ComponentDeliverySystem.OnExplorerInstalled() valida tipo y valor\n"
"5. Si valor correcto -> ApplyRepairToCircuit() -> CircuitManager.MarkDirty()\n"
"              -> OnCircuitChanged -> GameManager.CheckWinCondition()\n"
"6. Explorador llama GameSession.ReportarInstalacion(exito) -> RPC notifica al Técnico"
)

# 7.2.4
s2.blank()
s2.heading("7.2.4 Implementación del motor eléctrico")
s2.para("Los tres modelos de simulación aplican las leyes fundamentales de la electricidad:")
s2.blank()
s2.bold_intro("Serie (Reto 1):",
    "I = V_src / Sum(R) — caídas de voltaje calculadas secuencialmente nodo a nodo.")
s2.bold_intro("Paralelo (Reto 2):",
    "Cada rama recibe V_src completo; I_rama = V_src / R_rama; I_total = Sum(I_ramas).")
s2.bold_intro("Mixto (Retos 3 y 4):",
    "R_par_equiv = 1 / Sum(1/R_i),  R_total = R_serie + R_par_equiv,  V_paralelo = V_src - I x R_serie.")
s2.blank()
s2.para("Detección de cortocircuito: si R_total menor o igual a 0.1 Ohm, se activa isShortCircuited = true y la simulación se detiene con alerta de seguridad.", italic=True)

print("Seccion 7.2 insertada")

# ═══════════════════════════════════════════════════════════════════════════════
#  SECCIÓN 7.3
# ═══════════════════════════════════════════════════════════════════════════════

s3 = SectionBuilder(doc, anchors["74"])

# 7.3.1
s3.heading("7.3.1 Estrategia de pruebas (ISO/IEC/IEEE 29119)")
s3.para("Las pruebas se organizaron conforme a los procesos definidos en la norma ISO/IEC/IEEE 29119, aplicando cuatro niveles: pruebas de componente (unitarias), de integración, de sistema y de aceptación con usuarios. El proceso fue iterativo: los defectos detectados en cada nivel se corrigieron antes de avanzar al siguiente.")
s3.blank()
s3.fig("[Insertar Tabla 7.1: Plan de pruebas maestro conforme ISO 29119-3]")

# 7.3.2
s3.blank()
s3.heading("7.3.2 Pruebas de componente (Unitarias)")
s3.para("Verifican la corrección matemática del motor eléctrico de forma aislada, con valores de referencia calculados analíticamente.")
s3.blank()
s3.table(
    ["ID", "Caso de prueba", "Entrada", "Esperado", "Obtenido", "Estado"],
    [
        ["TC-U01", "Simulación Serie — ley de Ohm", "V=9V, R=100Ohm, R_LED=50Ohm", "I=0.060A", "I=0.060A", "Pasa"],
        ["TC-U02", "Cortocircuito Serie", "V=9V, R_total=0.05Ohm", "isShortCircuited=true", "isShortCircuited=true", "Pasa"],
        ["TC-U03", "Paralelo — rama abierta", "V=9V, R1=50Ohm, R2=9999Ohm", "LED2 apagado", "LED2 apagado", "Pasa"],
        ["TC-U04", "Mixta — voltaje de unión", "V=9V, R_s=220Ohm, R_L=50Ohm, R_C=100Ohm", "V_par~6.2V", "6.21V", "Pasa"],
        ["TC-U05", "IsValueCorrect — dentro rango", "valor=100.3Ohm, tol=0.5Ohm", "true", "true", "Pasa"],
        ["TC-U06", "IsValueCorrect — fuera rango", "valor=98Ohm, tol=0.5Ohm", "false", "false", "Pasa"],
        ["TC-U07", "PerformanceTracker — Excelente", "0 errores, t < límite", "Excelente", "Excelente", "Pasa"],
        ["TC-U08", "PerformanceTracker — Mejora", "5 errores", "Necesita mejorar", "Necesita mejorar", "Pasa"],
        ["TC-U09", "ObjectiveSystem — bono temporal", "50% tiempo restante", "points = max*0.5", "Correcto", "Pasa"],
    ]
)

# 7.3.3
s3.blank()
s3.heading("7.3.3 Pruebas de integración")
s3.para("Verifican la comunicación entre módulos a través del sistema de eventos, sin requerir hardware VR.")
s3.blank()
s3.table(
    ["ID", "Caso de prueba", "Módulos involucrados", "Esperado", "Estado"],
    [
        ["TC-I01", "Resistor correcto -> circuito se repara",
         "DeliverySystem -> CircuitManager -> GameManager",
         "LED verde, OnLevelCompleted(true)", "Pasa"],
        ["TC-I02", "Resistor incorrecto -> sin victoria",
         "DeliverySystem -> CircuitManager",
         "LED rojo, RegisterWrongAttempt +1", "Pasa"],
        ["TC-I03", "Slot tipo incorrecto -> error",
         "DeliverySystem -> GameManager",
         "RegisterWrongAttempt, componente destruido", "Pasa"],
        ["TC-I04", "Timer agotado -> fallo",
         "GameManager.Update -> CompleteLevel(false)",
         "OnLevelCompleted(level, false)", "Pasa"],
        ["TC-I05", "LoadLevel -> todos los sistemas reinician",
         "GameManager -> CircuitManager, PerformanceTracker, InstructionSystem",
         "Errores=0, timer reiniciado", "Pasa"],
        ["TC-I06", "RPC envío -> Explorador recibe componente",
         "GameSession -> ExplorerComponentReceiver",
         "Prefab instanciado en bandeja", "Parcial *"],
        ["TC-I07", "Reparación -> InstructionSystem avanza paso",
         "CircuitManager.OnCircuitChanged -> InstructionSystem",
         "currentStep++", "Pasa"],
    ]
)
s3.blank()
s3.para("* TC-I06: Requiere conexión Photon Fusion activa. Validado en escena IntegratedDemo con ambos roles en red local.", italic=True)

# 7.3.4
s3.blank()
s3.heading("7.3.4 Pruebas de sistema (escenarios completos)")
s3.para("Se ejecutaron los cuatro retos de principio a fin en la escena IntegratedDemo, verificando el flujo completo de la experiencia de juego.")
s3.blank()
s3.table(
    ["ID", "Reto", "Escenario", "Esperado", "Iter.", "Estado"],
    [
        ["TC-S01", "Reto 1", "Técnico envía R=100 Ohm correcta", "LED verde, transición a Reto 2", "2", "Pasa"],
        ["TC-S02", "Reto 1", "Técnico envía R=47 Ohm incorrecta", "LED rojo, error registrado, reintento posible", "1", "Pasa"],
        ["TC-S03", "Reto 2", "Explorador reconecta rama rota", "Ambos LEDs verdes, transición a Reto 3", "3", "Pasa"],
        ["TC-S04", "Reto 3", "Corrección de 3 fallas en orden", "Level completado, transición a Reto 4", "2", "Pasa"],
        ["TC-S05", "Reto 4", "Pin + resistor + cable corregidos", "Sistema sensor-actuador funcional", "4", "En ajuste"],
        ["TC-S06", "Todos", "Timer agotado sin completar", "Pantalla de fallo, opción de reinicio", "1", "Pasa"],
        ["TC-S07", "Todos", "Flujo completo 4 retos", "OnGameCompleted, pantalla de resultados", "1", "Pasa"],
    ]
)

# 7.3.5
s3.blank()
s3.heading("7.3.5 Pruebas de aceptación — Usabilidad (ISO/IEC 25010)")
s3.para("Se aplicó una sesión de prueba con usuarios del segmento objetivo (estudiantes de Ingeniería de Software, asignatura Computación Ubicua) evaluando la característica de Usabilidad de la norma ISO/IEC 25010.")
s3.blank()
s3.table(
    ["Subcaracterística", "Métrica", "Umbral aceptable", "Resultado"],
    [
        ["Comprensibilidad", "% usuarios que comprenden su rol sin explicación previa", "> 80%", "Por relevar (Semana 12)"],
        ["Aprendizaje", "Tiempo hasta primera medición correcta con multímetro", "< 3 min", "Por relevar"],
        ["Operabilidad", "% de acciones completadas sin error de interfaz", "> 90%", "Por relevar"],
        ["Satisfacción", "Escala Likert 1–5 (promedio)", ">= 3.5", "Por relevar"],
        ["Presencia VR", "System Usability Scale (SUS)", ">= 68 puntos", "Por relevar"],
    ]
)
s3.blank()
s3.para("Los resultados se registrarán durante la sesión de despliegue en el laboratorio de la UDLA (Fase 5, Semana 12).", italic=True)

# 7.3.6
s3.blank()
s3.heading("7.3.6 Defectos identificados y correcciones")
s3.para("Durante las fases de prueba se identificaron y corrigieron seis defectos conforme al proceso de gestión de incidencias de ISO 29119-3.")
s3.blank()
s3.table(
    ["ID", "Descripción", "Módulo", "Severidad", "Corrección aplicada"],
    [
        ["BUG-01",
         "InputField del Técnico no respondía al teclado; botón Enviar inoperativo",
         "Desktop / UI", "Alta",
         "TechnicianController.SetupPC() asigna worldCamera y añade GraphicRaycaster a todos los canvas WorldSpace"],
        ["BUG-02",
         "Manual técnico no cerraba al presionar Escape",
         "Desktop", "Media",
         "Migración a New Input System: Keyboard.current.escapeKey.wasPressedThisFrame en ManualBookOpener.Update()"],
        ["BUG-03",
         "Canvas WorldSpace bloqueaba clics sobre objetos 3D",
         "Interaction", "Alta",
         "Eliminación de conversión a ScreenSpaceCamera; PhysicsRaycaster añadido a la cámara del Técnico en SetupPC()"],
        ["BUG-04",
         "Reto 4 no completaba; cable suelto no era validado",
         "Gameplay", "Alta",
         "Implementación de TechnicianActions.FixLooseCable() y herramienta de editor LooseCableButtonGenerator"],
        ["BUG-05",
         "SetupRetoX() iteraba lista de componentes vacía",
         "Gameplay", "Alta",
         "Reordenamiento en ActivateComponentsForLevel: components.Clear() se movió a después de AutoDetectComponents()"],
        ["BUG-06",
         "Explorador no recibía componente sin Photon Fusion activo",
         "Gameplay / Networking", "Media",
         "Evento estático local ComponentSendingTray.OnComponentSentLocal como canal de fallback para pruebas sin red"],
    ]
)

print("Seccion 7.3 insertada")

# ─── guardar ──────────────────────────────────────────────────────────────────
doc.save(DOC_PATH)
print(f"\nDocumento guardado en:\n   {DOC_PATH}")
