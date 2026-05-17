# -*- coding: utf-8 -*-
"""
Genera diagramas C4 Nivel 3 (Componentes) para el Formato Capstone.docx
Produce dos PNG y los inserta en el documento Word.
"""

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
import matplotlib.patheffects as pe
import numpy as np
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUT_DIR  = r"C:\Users\holaq\Proyecto-TITA"
DOC_PATH = os.path.join(OUT_DIR, "Formato Capstone.docx")

# ── Paleta de colores C4 ──────────────────────────────────────────────────────
C_CONTAINER_BG   = "#ddeeff"   # azul claro — boundary del contenedor
C_CONTAINER_LINE = "#4477aa"
C_COMP_ELEC      = "#1168bd"   # azul oscuro — componentes Electrical
C_COMP_GAME      = "#2e7d32"   # verde oscuro — componentes Gameplay
C_COMP_NET       = "#6a1b9a"   # morado     — Networking
C_COMP_PLAYER    = "#e65100"   # naranja    — Player
C_COMP_UI        = "#00838f"   # teal       — UI / Desktop
C_COMP_TEXT      = "white"
C_ARROW          = "#555555"
C_EXT_BG         = "#cccccc"
C_EXT_LINE       = "#888888"


# ─────────────────────────────────────────────────────────────────────────────
#  Clase Box — rectángulo con título, subtítulo, descripción y color
# ─────────────────────────────────────────────────────────────────────────────
class Box:
    def __init__(self, ax, x, y, w, h, title, subtitle="", desc="",
                 color="#1168bd", text_color="white", alpha=1.0, lw=1.5):
        self.ax = ax
        self.x, self.y, self.w, self.h = x, y, w, h
        self.cx = x + w / 2
        self.cy = y + h / 2

        rect = FancyBboxPatch((x, y), w, h,
                              boxstyle="round,pad=0.02",
                              facecolor=color, edgecolor="white",
                              linewidth=lw, alpha=alpha, zorder=3)
        ax.add_patch(rect)

        # Título
        ty = y + h - 0.18
        ax.text(self.cx, ty, title,
                ha="center", va="top", fontsize=8, fontweight="bold",
                color=text_color, zorder=4, wrap=True,
                multialignment="center")

        # Subtítulo en cursiva
        if subtitle:
            ax.text(self.cx, ty - 0.19, f"[{subtitle}]",
                    ha="center", va="top", fontsize=6.5, fontstyle="italic",
                    color=text_color, zorder=4, alpha=0.9)

        # Descripción
        if desc:
            dy = ty - (0.38 if subtitle else 0.19)
            ax.text(self.cx, dy, desc,
                    ha="center", va="top", fontsize=5.8,
                    color=text_color, zorder=4,
                    wrap=True, multialignment="center",
                    linespacing=1.3)

    def top(self):    return (self.cx, self.y + self.h)
    def bottom(self): return (self.cx, self.y)
    def left(self):   return (self.x, self.cy)
    def right(self):  return (self.x + self.w, self.cy)


def arrow(ax, start, end, label="", color=C_ARROW, style="->"):
    ax.annotate("", xy=end, xytext=start,
                arrowprops=dict(arrowstyle=style, color=color,
                                lw=1.2, connectionstyle="arc3,rad=0.0"),
                zorder=5)
    if label:
        mx = (start[0] + end[0]) / 2
        my = (start[1] + end[1]) / 2
        ax.text(mx, my, label, ha="center", va="center",
                fontsize=5.5, color=color,
                bbox=dict(boxstyle="round,pad=0.15", fc="white", ec="none", alpha=0.85),
                zorder=6)


def container_boundary(ax, x, y, w, h, title, color=C_CONTAINER_LINE):
    rect = FancyBboxPatch((x, y), w, h,
                          boxstyle="round,pad=0.03",
                          facecolor=C_CONTAINER_BG, edgecolor=color,
                          linewidth=2, linestyle="--", alpha=0.35, zorder=1)
    ax.add_patch(rect)
    ax.text(x + 0.08, y + h - 0.06, title,
            ha="left", va="top", fontsize=7.5, fontweight="bold",
            color=color, zorder=2)


def legend_patch(color, label):
    return mpatches.Patch(facecolor=color, edgecolor="white", label=label)


# ═════════════════════════════════════════════════════════════════════════════
#  FIGURA 7.2 — Módulos Gameplay y Electrical
# ═════════════════════════════════════════════════════════════════════════════

fig, ax = plt.subplots(figsize=(16, 11))
ax.set_xlim(0, 16)
ax.set_ylim(0, 11)
ax.axis("off")
fig.patch.set_facecolor("white")

ax.text(8, 10.75, "Figura 7.2 — C4 Nivel 3: Componentes — Módulos Gameplay y Electrical",
        ha="center", va="top", fontsize=12, fontweight="bold", color="#222222")
ax.text(8, 10.45, "Sistema cliente Unity (ambos roles)",
        ha="center", va="top", fontsize=9, fontstyle="italic", color="#555555")

# Boundaries
container_boundary(ax, 0.3, 0.3, 15.4, 9.9, "«Container» Cliente Unity — Gameplay & Electrical")

# ── Bloque Electrical (izquierda) ──────────────────────────────────────────
container_boundary(ax, 0.6, 0.55, 4.8, 7.8, "«Módulo» Electrical", color=C_COMP_ELEC)

bCircuit = Box(ax, 0.85, 6.4, 4.3, 1.6,
               "CircuitManager",
               subtitle="Component",
               desc="Simula Serie / Paralelo / Mixto.\nDirty-flag 20 Hz.\nDispara OnCircuitChanged.",
               color=C_COMP_ELEC)

bElComp = Box(ax, 0.85, 4.6, 4.3, 1.5,
              "ElectricalComponent",
              subtitle="Abstract Component",
              desc="Resistor · LED · Capacitor\nVoltageSource · ArduinoPin\nGetResistance() / Calculate()",
              color=C_COMP_ELEC)

bNode = Box(ax, 0.85, 3.3, 2.0, 1.0,
            "ElectricalNode",
            subtitle="Component",
            desc="voltage · current",
            color=C_COMP_ELEC)

bPerf = Box(ax, 3.15, 3.3, 2.0, 1.0,
            "PerformanceTracker",
            subtitle="Component",
            desc="Errores · tiempo\nGetEvaluation()",
            color=C_COMP_ELEC)

bAnalyzer = Box(ax, 0.85, 1.9, 4.3, 1.1,
                "CircuitAnalyzer",
                subtitle="Component",
                desc="AnalyzeVoltage() · AnalyzeByLevel()\nFeedback educativo al Técnico",
                color=C_COMP_ELEC)

# ── Bloque Gameplay (centro-derecha) ──────────────────────────────────────
container_boundary(ax, 5.7, 0.55, 10.0, 7.8, "«Módulo» Gameplay", color=C_COMP_GAME)

bGM = Box(ax, 5.95, 7.5, 4.2, 1.55,
          "GameManager",
          subtitle="Component",
          desc="Máquina de estados (LevelType).\nOrquesta 4 retos secuenciales.\nOnLevelLoaded / Completed / Timer.",
          color=C_COMP_GAME)

bInstr = Box(ax, 10.45, 7.5, 4.95, 1.55,
             "InstructionSystem",
             subtitle="Component",
             desc="Instrucciones paso a paso.\nValidación automática por evento.\nGetCurrentInstruction()",
             color=C_COMP_GAME)

bDeliv = Box(ax, 5.95, 5.6, 4.2, 1.6,
             "ComponentDeliverySystem",
             subtitle="Component",
             desc="Envío-Instalación-Validación.\nSendResistor/LED/Capacitor/Pin.\nOnRepairValidated event.",
             color=C_COMP_GAME)

bSlot = Box(ax, 10.45, 5.6, 2.2, 1.6,
            "ComponentSlot",
            subtitle="Component",
            desc="Valida tipo\ny valor del\ncomponente.",
            color=C_COMP_GAME)

bReceiver = Box(ax, 12.95, 5.6, 2.2, 1.6,
                "ExplorerComponent\nReceiver",
                subtitle="Component",
                desc="Instancia prefab\nen bandeja\ndel Explorador.",
                color=C_COMP_GAME)

bObj = Box(ax, 5.95, 3.95, 2.8, 1.4,
           "ObjectiveSystem",
           subtitle="Component",
           desc="Puntuación + bono\ntemporal por reto.\nOnSessionEnded.",
           color=C_COMP_GAME)

bDiag = Box(ax, 9.05, 3.95, 3.1, 1.4,
            "DiagnosticSystem",
            subtitle="Component",
            desc="Feedback al Técnico\nsobre el estado\ndel circuito.",
            color=C_COMP_GAME)

bWire = Box(ax, 12.45, 3.95, 3.0, 1.4,
            "CircuitWireRenderer",
            subtitle="Component",
            desc="Renderiza cables\ny flujo visual\ndel circuito.",
            color=C_COMP_GAME)

bTechAct = Box(ax, 5.95, 2.2, 4.2, 1.45,
               "TechnicianActions",
               subtitle="Component",
               desc="Acciones del Técnico:\nHasSelectedResistor()\nFixLooseCable()",
               color=C_COMP_GAME)

bTechMan = Box(ax, 10.45, 2.2, 5.0, 1.45,
               "TechnicianManual",
               subtitle="Component",
               desc="Manual técnico interactivo.\nDiagramas y tablas de referencia\npara diagnóstico de fallas.",
               color=C_COMP_GAME)

bZone = Box(ax, 5.95, 0.75, 4.2, 1.15,
            "ZoneProximityScaler",
            subtitle="Component",
            desc="Escala objetos según\nproximidad del Explorador.",
            color=C_COMP_GAME)

bLvl = Box(ax, 10.45, 0.75, 5.0, 1.15,
           "LevelType · ChallengeTag",
           subtitle="Enum / Component",
           desc="OhmLaw · Parallel · Mixed · Arduino\nEtiquetas de reto para cada zona.",
           color=C_COMP_GAME)

# ── Flechas Electrical ──────────────────────────────────────────────────────
arrow(ax, bCircuit.bottom(), bElComp.top(), "usa")
arrow(ax, bElComp.bottom(), (bElComp.cx, bNode.y + bNode.h), "tiene nodos")
arrow(ax, bCircuit.right(), (bCircuit.x + bCircuit.w, bPerf.cy), "")

# ── Flechas Gameplay ────────────────────────────────────────────────────────
arrow(ax, bCircuit.right(), bGM.left(), "OnCircuitChanged\n→ CheckWinCondition")
arrow(ax, bGM.bottom(), bDeliv.top(), "RegisterRepair\n/ WrongAttempt")
arrow(ax, bGM.right(), bInstr.left(), "OnLevelLoaded\nBuildInstructions")
arrow(ax, bDeliv.right(), bSlot.left(), "instala en")
arrow(ax, bSlot.right(), bReceiver.left(), "recibe de\nExplorador")
arrow(ax, bGM.bottom(), (bGM.cx, bObj.y + bObj.h), "OnLevelCompleted")
arrow(ax, bInstr.bottom(), (bInstr.cx, bDiag.y + bDiag.h), "estado circuito")
arrow(ax, bGM.bottom(), bTechAct.top(), "consulta\nacciones")
arrow(ax, bTechAct.right(), bTechMan.left(), "usa")

# Leyenda
legend_handles = [
    legend_patch(C_COMP_ELEC, "Módulo Electrical"),
    legend_patch(C_COMP_GAME, "Módulo Gameplay"),
]
ax.legend(handles=legend_handles, loc="lower right",
          fontsize=7.5, framealpha=0.9,
          bbox_to_anchor=(1.0, 0.0))

plt.tight_layout(rect=[0, 0, 1, 0.97])
path72 = os.path.join(OUT_DIR, "fig72_c4_gameplay_electrical.png")
plt.savefig(path72, dpi=180, bbox_inches="tight", facecolor="white")
plt.close()
print(f"Figura 7.2 guardada: {path72}")


# ═════════════════════════════════════════════════════════════════════════════
#  FIGURA 7.3 — Módulos Networking y Player
# ═════════════════════════════════════════════════════════════════════════════

fig, ax = plt.subplots(figsize=(16, 10))
ax.set_xlim(0, 16)
ax.set_ylim(0, 10)
ax.axis("off")
fig.patch.set_facecolor("white")

ax.text(8, 9.75, "Figura 7.3 — C4 Nivel 3: Componentes — Módulos Networking y Player",
        ha="center", va="top", fontsize=12, fontweight="bold", color="#222222")
ax.text(8, 9.45, "Sistema cliente Unity (ambos roles)",
        ha="center", va="top", fontsize=9, fontstyle="italic", color="#555555")

container_boundary(ax, 0.3, 0.3, 15.4, 8.9, "«Container» Cliente Unity — Networking, Player, UI & Desktop")

# ── Módulo Networking ──────────────────────────────────────────────────────
container_boundary(ax, 0.6, 5.4, 6.9, 3.4, "«Módulo» Networking", color=C_COMP_NET)

bConn = Box(ax, 0.85, 7.1, 6.4, 1.5,
            "ConnectionManager",
            subtitle="Component · INetworkRunnerCallbacks",
            desc="Inicia sesión Photon Fusion según rol.\nHost (Técnico) / Client (Explorador).\nAuto-connect configurable desde Inspector.",
            color=C_COMP_NET)

bSession = Box(ax, 0.85, 5.6, 6.4, 1.3,
               "GameSession",
               subtitle="NetworkBehaviour",
               desc="[Networked] RetoActual · HayComponentePendiente · TipoComponente · ValorComponente\n"
                    "RPC_EnviarComponente · RPC_ComponenteInstalado · RPC_CambiarReto",
               color=C_COMP_NET)

# ── Módulo Player ──────────────────────────────────────────────────────────
container_boundary(ax, 0.6, 0.5, 6.9, 4.65, "«Módulo» Player", color=C_COMP_PLAYER)

bPlayerCtrl = Box(ax, 0.85, 3.55, 6.4, 1.4,
                  "PlayerController",
                  subtitle="Component · CharacterController",
                  desc="Locomoción KAT VR (SDK) con fallback joystick Meta Quest.\n"
                       "CalibrateOrientation() · FreezeMovement().\n"
                       "Sincroniza XR Origin con CharacterController en LateUpdate.",
                  color=C_COMP_PLAYER)

bHaptic = Box(ax, 0.85, 2.35, 3.0, 0.95,
              "HapticFeedback",
              subtitle="Component",
              desc="Pulsos al chaleco háptico\nproporcionales a la corriente\ndel circuito simulado.",
              color=C_COMP_PLAYER)

bInteraction = Box(ax, 4.1, 2.35, 3.1, 0.95,
                   "PlayerInteraction",
                   subtitle="Component",
                   desc="Detección de objetos XRI.\nGestiona estado de\ninteracción activa.",
                   color=C_COMP_PLAYER)

bTechCtrl = Box(ax, 0.85, 1.2, 3.0, 0.9,
                "TechnicianController",
                subtitle="Component",
                desc="Interacción ratón en PC.\nPhysicsRaycaster.\nSetupPC().",
                color=C_COMP_PLAYER)

bToolbox = Box(ax, 4.1, 1.2, 3.1, 0.9,
               "ToolboxController",
               subtitle="Component",
               desc="Gestiona herramientas\ndel Técnico\n(multímetro, manual).",
               color=C_COMP_PLAYER)

bAvatar = Box(ax, 0.85, 0.55, 6.4, 0.45,
              "ExplorerAvatar  ·  MouseGrabSimulator (Core)",
              subtitle="Components",
              desc="",
              color=C_COMP_PLAYER)

# ── Módulo UI / Desktop ────────────────────────────────────────────────────
container_boundary(ax, 7.8, 5.4, 7.9, 3.4, "«Módulo» UI", color=C_COMP_UI)

bHUD_T = Box(ax, 8.05, 7.1, 3.6, 1.5,
             "TechnicianHUDController",
             subtitle="Component",
             desc="Instrucciones del paso actual.\nTemporizador · Puntaje.\nIndicadores de falla del circuito.",
             color=C_COMP_UI)

bHUD_E = Box(ax, 11.95, 7.1, 3.5, 1.5,
             "ExplorerHUD / PlayerFeedbackUI",
             subtitle="Components",
             desc="HUD del Explorador VR.\nFeedback visual de estado.\nMensajes de resultado.",
             color=C_COMP_UI)

bMultimeterUI = Box(ax, 8.05, 5.6, 3.6, 1.3,
                    "MultimeterUI",
                    subtitle="Component",
                    desc="Muestra voltaje medido\npor el Explorador.\nActualización en tiempo real.",
                    color=C_COMP_UI)

bDiagramPanel = Box(ax, 11.95, 5.6, 3.5, 1.3,
                    "CircuitDiagramPanel",
                    subtitle="Component",
                    desc="Esquema del circuito activo\ncon indicadores de falla\nen tiempo real.",
                    color=C_COMP_UI)

# ── Módulo Desktop ─────────────────────────────────────────────────────────
container_boundary(ax, 7.8, 0.5, 7.9, 4.65, "«Módulo» Desktop", color="#795548")

bWorkstation = Box(ax, 8.05, 3.55, 3.6, 1.4,
                   "TechnicianWorkstation",
                   subtitle="Component",
                   desc="Interfaz estación de trabajo.\nSelección y envío de componentes.\nGestiona manuales técnicos.",
                   color="#795548")

bTray = Box(ax, 11.95, 3.55, 3.5, 1.4,
            "ComponentSendingTray",
            subtitle="Component",
            desc="Materializa componente\nseleccionado.\nEnvía por red (Photon / local).",
            color="#795548")

bDeskComp = Box(ax, 8.05, 2.35, 3.6, 0.95,
                "DeskComponent",
                subtitle="Component",
                desc="Objetos de escritorio\ninteractivos con glow\nal seleccionar.",
                color="#795548")

bManual = Box(ax, 11.95, 2.35, 3.5, 0.95,
              "ManualBookOpener\n/ TechnicianManualDisplay",
              subtitle="Components",
              desc="Abre/cierra el manual técnico.\nMuestra diagramas y tablas.",
              color="#795548")

bUIBtn = Box(ax, 8.05, 1.2, 7.4, 0.9,
             "UIButtonController",
             subtitle="Component",
             desc="Controla botones de la UI del Técnico: enviar componente, seleccionar modo, confirmar acción.",
             color="#795548")

# ── Sistemas externos ──────────────────────────────────────────────────────
bPhoton = Box(ax, 0.6, 9.4, 6.9, 0.38,
              "Photon Fusion Cloud  [Servicio externo de red]",
              color=C_EXT_BG, text_color="#222222", lw=1)
bKAT = Box(ax, 7.8, 9.4, 3.8, 0.38,
           "KAT VR SDK  [Hardware externo]",
           color=C_EXT_BG, text_color="#222222", lw=1)
bXRI = Box(ax, 11.85, 9.4, 3.85, 0.38,
           "XR Interaction Toolkit  [Paquete Unity]",
           color=C_EXT_BG, text_color="#222222", lw=1)

# ── Flechas ────────────────────────────────────────────────────────────────
# Networking
arrow(ax, bConn.bottom(), bSession.top(), "crea y gestiona")
arrow(ax, (bSession.x + bSession.w, bSession.cy),
          (bHUD_T.x, bSession.cy), "OnComponenteRecibido\nOnRetoChanged")

# Player
arrow(ax, bPlayerCtrl.bottom(), bHaptic.top(), "feedback")
arrow(ax, bPlayerCtrl.bottom(), bInteraction.top(), "interacción")
arrow(ax, bHaptic.bottom(), bTechCtrl.top(), "")
arrow(ax, bInteraction.bottom(), bToolbox.top(), "")

# Networking → Player
arrow(ax, bSession.bottom(), bPlayerCtrl.top(), "Networked state\n→ avatar sync")

# UI ← GameManager (externo al recuadro, representado con label)
arrow(ax, (bHUD_T.cx, bHUD_T.y),
          (bHUD_T.cx, bHUD_T.y - 0.3), "← OnLevelLoaded\n← OnTimerTick")

# Desktop → Networking
arrow(ax, bTray.left(), (bSession.x + bSession.w, bTray.cy),
      "EnviarComponente()\n→ RPC Fusion")

# Leyenda
legend_handles = [
    legend_patch(C_COMP_NET,    "Módulo Networking"),
    legend_patch(C_COMP_PLAYER, "Módulo Player"),
    legend_patch(C_COMP_UI,     "Módulo UI"),
    legend_patch("#795548",     "Módulo Desktop"),
    legend_patch(C_EXT_BG,      "Sistema externo"),
]
ax.legend(handles=legend_handles, loc="lower right",
          fontsize=7, framealpha=0.9,
          bbox_to_anchor=(1.0, 0.0))

plt.tight_layout(rect=[0, 0, 1, 0.97])
path73 = os.path.join(OUT_DIR, "fig73_c4_networking_player.png")
plt.savefig(path73, dpi=180, bbox_inches="tight", facecolor="white")
plt.close()
print(f"Figura 7.3 guardada: {path73}")


# ═════════════════════════════════════════════════════════════════════════════
#  INSERTAR FIGURAS EN EL WORD
# ═════════════════════════════════════════════════════════════════════════════

def insert_image_replacing_placeholder(doc, placeholder_text, image_path, caption):
    """Reemplaza un párrafo placeholder con la imagen y su pie de figura."""
    for i, p in enumerate(doc.paragraphs):
        if placeholder_text in p.text:
            # Limpiar el párrafo placeholder
            for run in p.runs:
                run.text = ""
            # Insertar imagen como run en ese párrafo
            run = p.add_run()
            run.add_picture(image_path, width=Inches(6.2))
            p.alignment = 1  # CENTER

            # Insertar párrafo de pie de figura justo después
            # Crear nuevo párrafo después del actual
            new_p_el = OxmlElement("w:p")
            p._element.addnext(new_p_el)
            from docx.text.paragraph import Paragraph as DocPara
            cap_para = DocPara(new_p_el, p._element.getparent())
            cap_run = cap_para.add_run(caption)
            cap_run.italic = True
            cap_run.font.size = Pt(9)
            cap_para.alignment = 1  # CENTER
            print(f"  Imagen insertada para: {placeholder_text}")
            return True
    print(f"  AVISO: No se encontro placeholder: {placeholder_text}")
    return False

doc = Document(DOC_PATH)
insert_image_replacing_placeholder(
    doc,
    "[Insertar Figura 7.2: Diagrama C4 Nivel 3 — Módulos Gameplay y Electrical]",
    path72,
    "Figura 7.2. Diagrama C4 Nivel 3 — Módulos Gameplay y Electrical."
)
insert_image_replacing_placeholder(
    doc,
    "[Insertar Figura 7.3: Diagrama C4 Nivel 3 — Módulos Networking y Player]",
    path73,
    "Figura 7.3. Diagrama C4 Nivel 3 — Módulos Networking y Player."
)
doc.save(DOC_PATH)
print(f"\nDocumento actualizado: {DOC_PATH}")
